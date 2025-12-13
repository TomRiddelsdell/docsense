"""
Tests for Word converter exception handling.

Tests specific exception scenarios for Word document conversion.
"""
import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from src.infrastructure.converters.word_converter import WordConverter
from src.infrastructure.converters.exceptions import (
    InvalidFileFormatError,
    PasswordProtectedError,
    FileTooLargeError,
    FileNotReadableError,
)


class TestWordConverterExceptions:
    """Tests for Word converter exception handling."""

    def setup_method(self):
        """Set up test fixtures."""
        self.converter = WordConverter()

    def test_file_not_found_returns_error_result(self):
        """Test that FileNotFoundError is handled gracefully."""
        result = self.converter.convert(Path("/nonexistent/file.docx"))

        assert result.success is False
        assert len(result.errors) > 0
        assert "not found" in result.errors[0].lower()

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_package_not_found_error_raises_invalid_format(self, mock_document):
        """Test that PackageNotFoundError raises InvalidFileFormatError."""
        from docx.opc.exceptions import PackageNotFoundError

        # Mock package not found error
        mock_document.side_effect = PackageNotFoundError("Not a valid DOCX package")

        docx_content = b"PK\x03\x04 invalid"

        with pytest.raises(InvalidFileFormatError) as exc_info:
            self.converter.convert_from_bytes(docx_content, "invalid.docx")

        assert exc_info.value.details.get('format') == 'DOCX'
        assert exc_info.value.details.get('filename') == "invalid.docx"

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_invalid_xml_error_raises_invalid_format(self, mock_document):
        """Test that InvalidXmlError raises InvalidFileFormatError."""
        from docx.oxml.exceptions import InvalidXmlError

        # Mock OXML package read error
        mock_document.side_effect = InvalidXmlError("Corrupted XML")

        docx_content = b"PK\x03\x04 corrupted"

        with pytest.raises(InvalidFileFormatError) as exc_info:
            self.converter.convert_from_bytes(docx_content, "corrupted.docx")

        assert exc_info.value.details.get('format') == 'DOCX'
        assert "corrupted" in str(exc_info.value).lower() or "malformed" in str(exc_info.value).lower()

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_memory_error_raises_file_too_large(self, mock_document):
        """Test that MemoryError raises FileTooLargeError with size details."""
        # Mock memory error
        mock_document.side_effect = MemoryError("Out of memory")

        # Large DOCX content (simulated)
        docx_content = b"PK\x03\x04" + b"x" * (30 * 1024 * 1024)  # 30MB

        with pytest.raises(FileTooLargeError) as exc_info:
            self.converter.convert_from_bytes(docx_content, "large.docx")

        assert exc_info.value.details.get('filename') == "large.docx"
        assert 'size_mb' in exc_info.value.details
        # Size should be approximately 30 MB
        size_mb = exc_info.value.details['size_mb']
        assert 29 < size_mb < 31

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_os_error_with_password_raises_password_protected(self, mock_document):
        """Test that OSError mentioning password raises PasswordProtectedError."""
        # Mock OS error with password message
        mock_document.side_effect = OSError("file is password-protected")

        docx_content = b"PK\x03\x04"

        with pytest.raises(PasswordProtectedError) as exc_info:
            self.converter.convert_from_bytes(docx_content, "protected.docx")

        assert exc_info.value.details.get('filename') == "protected.docx"

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_io_error_with_encrypted_raises_password_protected(self, mock_document):
        """Test that IOError mentioning encryption raises PasswordProtectedError."""
        # Mock IO error with encryption message
        mock_document.side_effect = IOError("encrypted document")

        docx_content = b"PK\x03\x04"

        with pytest.raises(PasswordProtectedError):
            self.converter.convert_from_bytes(docx_content, "encrypted.docx")

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_os_error_without_password_raises_invalid_format(self, mock_document):
        """Test that OSError without password/encryption raises InvalidFileFormatError."""
        # Mock OS error without password message
        mock_document.side_effect = OSError("Permission denied")

        docx_content = b"PK\x03\x04"

        with pytest.raises(InvalidFileFormatError) as exc_info:
            self.converter.convert_from_bytes(docx_content, "denied.docx")

        assert exc_info.value.details.get('filename') == "denied.docx"

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_large_docx_returns_warning(self, mock_document):
        """Test that large DOCX files (>20MB) return size warning."""
        # Mock valid document
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = []
        mock_doc.core_properties.title = "Test"
        mock_document.return_value = mock_doc

        # 25MB DOCX
        docx_content = b"PK\x03\x04" + b"x" * (25 * 1024 * 1024)

        result = self.converter.convert_from_bytes(docx_content, "large.docx")

        # Should succeed but with warning
        assert result.success is True
        assert len(result.warnings) > 0
        # Check for size warning
        assert any("MB" in w or "large" in w.lower() for w in result.warnings)

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_unexpected_error_is_reraised(self, mock_document):
        """Test that unexpected errors are re-raised with full context."""
        # Mock unexpected error
        mock_document.side_effect = ValueError("Unexpected error")

        docx_content = b"PK\x03\x04"

        with pytest.raises(ValueError, match="Unexpected error"):
            self.converter.convert_from_bytes(docx_content, "test.docx")


class TestWordTableExtractionExceptions:
    """Tests for Word table extraction exception handling."""

    def setup_method(self):
        """Set up test fixtures."""
        self.converter = WordConverter()

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_table_index_error_continues_conversion(self, mock_document):
        """Test that IndexError during table extraction doesn't fail conversion."""
        # Mock document with malformed table
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.core_properties.title = "Test"

        # Mock table that raises IndexError
        mock_table = MagicMock()
        mock_table.rows.__iter__.side_effect = IndexError("Invalid table structure")
        mock_doc.tables = [mock_table]

        mock_document.return_value = mock_doc

        docx_content = b"PK\x03\x04"

        result = self.converter.convert_from_bytes(docx_content, "test.docx")

        # Should succeed despite table error
        assert result.success is True
        # Should have warning about table extraction
        assert len(result.warnings) > 0
        assert any("table" in w.lower() for w in result.warnings)

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_table_attribute_error_continues_conversion(self, mock_document):
        """Test that AttributeError during table extraction doesn't fail conversion."""
        # Mock document with malformed table
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.core_properties.title = "Test"

        # Mock table that raises AttributeError
        mock_table = MagicMock()
        del mock_table.rows  # Remove rows attribute
        type(mock_table).rows = property(lambda self: (_ for _ in ()).throw(AttributeError("Missing attribute")))
        mock_doc.tables = [mock_table]

        mock_document.return_value = mock_doc

        docx_content = b"PK\x03\x04"

        result = self.converter.convert_from_bytes(docx_content, "test.docx")

        # Should succeed despite table error
        assert result.success is True
        # Should have warning
        assert len(result.warnings) > 0

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_partial_table_success(self, mock_document):
        """Test that some tables can succeed while others fail."""
        # Mock document with mixed tables
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.core_properties.title = "Test"

        # First table succeeds
        mock_good_table = MagicMock()
        mock_row1 = MagicMock()
        mock_cell1 = MagicMock()
        mock_cell1.text = "Header"
        mock_row1.cells = [mock_cell1]
        mock_good_table.rows = [mock_row1]

        # Second table fails
        mock_bad_table = MagicMock()
        mock_bad_table.rows.__iter__.side_effect = IndexError("Invalid")

        mock_doc.tables = [mock_good_table, mock_bad_table]
        mock_document.return_value = mock_doc

        docx_content = b"PK\x03\x04"

        result = self.converter.convert_from_bytes(docx_content, "test.docx")

        # Should succeed with partial results
        assert result.success is True
        # Should have some table content from good table
        assert "Header" in result.markdown_content or len(result.warnings) > 0

    @patch('src.infrastructure.converters.word_converter.Document')
    def test_generic_table_exception_continues(self, mock_document):
        """Test that generic exceptions during table processing are handled."""
        # Mock document with table that raises generic exception
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.core_properties.title = "Test"

        # Mock table that raises generic exception
        mock_table = MagicMock()
        mock_table.rows.__iter__.side_effect = RuntimeError("Unexpected error")
        mock_doc.tables = [mock_table]

        mock_document.return_value = mock_doc

        docx_content = b"PK\x03\x04"

        result = self.converter.convert_from_bytes(docx_content, "test.docx")

        # Should succeed despite table error
        assert result.success is True
        # Should have warning
        assert len(result.warnings) > 0
