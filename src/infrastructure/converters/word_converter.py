from pathlib import Path
from io import BytesIO
import logging

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.exceptions import InvalidXmlError

from .base import (
    DocumentConverter,
    ConversionResult,
    DocumentMetadata,
    DocumentFormat,
)
from .exceptions import (
    InvalidFileFormatError,
    FileTooLargeError,
    PasswordProtectedError,
)

logger = logging.getLogger(__name__)


class WordConverter(DocumentConverter):
    
    @property
    def supported_extensions(self) -> list[str]:
        return ["docx", "doc"]
    
    def convert(self, file_path: Path) -> ConversionResult:
        try:
            with open(file_path, 'rb') as f:
                return self.convert_from_bytes(f.read(), file_path.name)
        except FileNotFoundError:
            return ConversionResult(
                success=False,
                markdown_content="",
                sections=[],
                metadata=DocumentMetadata(original_format=DocumentFormat.WORD),
                errors=[f"File not found: {file_path}"]
            )
    
    def convert_from_bytes(self, content: bytes, filename: str) -> ConversionResult:
        errors = []
        warnings = []

        # Check file size (20 MB limit for Word documents)
        size_mb = len(content) / (1024 * 1024)
        if size_mb > 20:
            logger.warning(f"Large Word file: {size_mb:.1f} MB")
            warnings.append(f"Large file ({size_mb:.1f} MB) may take longer to process")

        try:
            doc = Document(BytesIO(content))

        except PackageNotFoundError as e:
            logger.error(f"Invalid Word document package: {filename}: {e}")
            raise InvalidFileFormatError(
                "File is not a valid Word document (DOCX)",
                details={'format': 'DOCX', 'filename': filename, 'error': str(e)}
            )

        except InvalidXmlError as e:
            # OXML package read errors - usually corruption or invalid XML
            logger.error(f"Corrupted Word document: {filename}: {e}")
            raise InvalidFileFormatError(
                "Word document is corrupted or malformed",
                details={'format': 'DOCX', 'filename': filename, 'error': str(e)}
            )

        except MemoryError:
            logger.critical(f"Out of memory processing Word document: {filename} ({size_mb:.1f} MB)")
            raise FileTooLargeError(
                f"File too large to process ({size_mb:.1f} MB)",
                details={'size_mb': size_mb, 'limit_mb': 20, 'filename': filename}
            )

        except (OSError, IOError) as e:
            # File access errors
            error_str = str(e).lower()
            if 'password' in error_str or 'encrypted' in error_str:
                logger.error(f"Password-protected Word document: {filename}")
                raise PasswordProtectedError(
                    "Document is password protected",
                    details={'filename': filename}
                )
            else:
                logger.error(f"Cannot read Word document: {filename}: {e}")
                raise InvalidFileFormatError(
                    "Cannot read Word document",
                    details={'format': 'DOCX', 'filename': filename, 'error': str(e)}
                )

        except Exception as e:
            # Log unexpected errors with full stack trace
            logger.exception(f"Unexpected error opening Word document {filename}: {e}")
            raise
        
        markdown_lines = []
        
        for para in doc.paragraphs:
            style_name: str = para.style.name if para.style and para.style.name else ""
            text = para.text.strip()
            
            if not text:
                markdown_lines.append("")
                continue
            
            if "Heading 1" in style_name:
                markdown_lines.append(f"# {text}")
            elif "Heading 2" in style_name:
                markdown_lines.append(f"## {text}")
            elif "Heading 3" in style_name:
                markdown_lines.append(f"### {text}")
            elif "Heading 4" in style_name:
                markdown_lines.append(f"#### {text}")
            elif "Heading 5" in style_name:
                markdown_lines.append(f"##### {text}")
            elif "Heading 6" in style_name:
                markdown_lines.append(f"###### {text}")
            elif "List" in style_name:
                markdown_lines.append(f"- {text}")
            elif "Code" in style_name:
                markdown_lines.append(f"```\n{text}\n```")
            else:
                markdown_lines.append(text)
            
            markdown_lines.append("")
        
        for table_idx, table in enumerate(doc.tables, 1):
            try:
                table_md = self._convert_table(table)
                markdown_lines.append(table_md)
                markdown_lines.append("")
            except (IndexError, AttributeError) as e:
                logger.warning(f"Skipping malformed table {table_idx}: {e}")
                warnings.append(f"Could not convert table {table_idx} (malformed structure)")
            except Exception as e:
                logger.warning(f"Error converting table {table_idx}: {e}")
                warnings.append(f"Could not convert table {table_idx}")
        
        markdown_content = '\n'.join(markdown_lines)
        
        core_props = doc.core_properties
        metadata = DocumentMetadata(
            title=core_props.title or filename,
            author=core_props.author,
            created_date=str(core_props.created) if core_props.created else None,
            modified_date=str(core_props.modified) if core_props.modified else None,
            page_count=len(doc.sections),
            word_count=self._count_words(markdown_content),
            original_format=DocumentFormat.WORD
        )
        
        sections = self._extract_sections(markdown_content)
        
        return ConversionResult(
            success=True,
            markdown_content=markdown_content,
            sections=sections,
            metadata=metadata,
            errors=errors,
            warnings=warnings
        )
    
    def _convert_table(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
        
        if len(rows) >= 1:
            header_sep = '| ' + ' | '.join(['---'] * len(table.rows[0].cells)) + ' |'
            rows.insert(1, header_sep)
        
        return '\n'.join(rows)
