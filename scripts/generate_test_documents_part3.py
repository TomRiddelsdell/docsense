#!/usr/bin/env python3
"""
Generate final 10 test documents (Part 3 of 3).
Documents 11-20 covering full spectrum from perfect to severely flawed.
"""

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


OUTPUT_DIR = Path("data/test_documents")


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text, style='Normal'):
    return doc.add_paragraph(text, style=style)


def add_formula(doc, formula_text):
    para = doc.add_paragraph()
    run = para.add_run(formula_text)
    run.italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return para


def save_document(doc, filename, metadata):
    doc_path = OUTPUT_DIR / f"{filename}.docx"
    json_path = OUTPUT_DIR / f"{filename}.json"
    doc.save(doc_path)
    with open(json_path, 'w') as f:
        json.dump(metadata, f, indent=2)
    print(f"Created: {doc_path.name} and {json_path.name}")


# Document 11: Nearly perfect, fully compliant
def create_doc_11_perfect_compliance():
    """Exemplary document with complete specifications"""
    doc = Document()
    
    add_heading(doc, "Russell 2000 Equal Weight Index - Methodology Document", 1)
    
    add_heading(doc, "1. Index Overview", 2)
    add_paragraph(doc, """
The Russell 2000 Equal Weight Index provides equal-weighted exposure to the 2000 smallest 
companies in the Russell 3000 Index, as defined by FTSE Russell. Each constituent receives 
a weight of 0.05% (1/2000) at each quarterly rebalancing.
""")
    
    add_heading(doc, "2. Eligibility Criteria", 2)
    add_paragraph(doc, """
Constituents must:
- Be included in the Russell 3000 Index as of the most recent Russell Reconstitution (last Friday in June)
- Have minimum average daily trading volume of $1,000,000 over the past 3 months
- Have minimum market capitalization of $200 million as measured on the reconstitution date
- Be incorporated and headquartered in the United States
- Trade on NYSE, NASDAQ, or NYSE American exchanges
""")
    
    add_heading(doc, "3. Data Sources", 2)
    add_paragraph(doc, """
- Index constituent list: FTSE Russell Index Data Feed, updated daily
- Pricing data: FTSE Russell Official Closing Prices, 16:00 ET
- Corporate actions: FTSE Russell Corporate Actions Data Feed, real-time
- Trading volume: Consolidated Tape, sourced from FTSE Russell
""")
    
    add_heading(doc, "4. Rebalancing", 2)
    add_paragraph(doc, """
Quarterly rebalancing on third Friday of March, June, September, and December.
Weights reset to equal weight (0.05% per constituent) using closing prices from the 
rebalancing date. Changes effective on market open the following Monday.
""")
    
    add_heading(doc, "5. Risk Disclosures", 2)
    add_paragraph(doc, """
Small-cap concentration risk: 100% exposure to small-capitalization stocks which historically 
exhibit higher volatility than large-caps. Equal weighting may result in tracking error vs 
market-cap weighted indices. Quarterly rebalancing generates transaction costs.
""")
    
    metadata = {
        "document_id": "doc_11_perfect_compliance",
        "title": "Russell 2000 Equal Weight Index - Methodology Document",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "low",
        "issue_count": 0,
        "issues": [],
        "self_containment_score": 0.98,
        "implementability_score": 0.98,
        "notes": "Exemplary document with complete specifications - ideal reference example"
    }
    
    save_document(doc, "doc_11_perfect_compliance", metadata)


# Document 12: Formula precision issues
def create_doc_12_formula_precision():
    """Document with imprecise mathematical specifications"""
    doc = Document()
    
    add_heading(doc, "Quantitative Mean Reversion Strategy", 1)
    
    add_heading(doc, "1. Signal Calculation", 2)
    add_paragraph(doc, """
Calculate z-score for each stock using historical prices.
""")
    add_formula(doc, "z = (current_price - mean_price) / std_dev")
    
    add_heading(doc, "2. Trading Rules", 2)
    add_paragraph(doc, """
Buy when z-score drops below threshold.
Sell when z-score returns to zero or exceeds upper threshold.
Position size based on signal strength.
""")
    
    add_heading(doc, "3. Risk Controls", 2)
    add_paragraph(doc, """
Maximum position size capped at reasonable level.
Stop loss triggered if loss exceeds limit.
""")
    
    metadata = {
        "document_id": "doc_12_formula_precision",
        "title": "Quantitative Mean Reversion Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 8,
        "issues": [
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Historical lookback period not specified",
                "description": "'historical prices' used in z-score but lookback period not defined",
                "location": "Section 1, formula",
                "original_text": "z = (current_price - mean_price) / std_dev",
                "suggested_fix": "Specify: 'mean_price and std_dev calculated using 60 trading days of daily closing prices'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Entry threshold not specified",
                "description": "'drops below threshold' but threshold value not defined",
                "location": "Section 2 - Trading Rules",
                "suggested_fix": "Specify: 'Buy when z-score < -2.0'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Exit threshold not specified",
                "description": "'exceeds upper threshold' but value not defined",
                "location": "Section 2 - Trading Rules",
                "suggested_fix": "Specify: 'Sell when z-score > +2.0'",
                "confidence": 0.95
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Position sizing formula not provided",
                "description": "'Position size based on signal strength' but no formula given",
                "location": "Section 2 - Trading Rules",
                "suggested_fix": "Specify: 'Position size = Base_Size × min(|z-score| / 2, 1.5)' or other explicit formula",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Maximum position size not defined",
                "description": "'reasonable level' is subjective",
                "location": "Section 3 - Risk Controls",
                "suggested_fix": "Specify: 'Maximum 5% of portfolio value per position'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Stop loss level not specified",
                "description": "'exceeds limit' but limit not defined",
                "location": "Section 3 - Risk Controls",
                "suggested_fix": "Specify: 'Stop loss at 3% loss from entry price'",
                "confidence": 0.90
            },
            {
                "severity": "high",
                "category": "incomplete_formula",
                "title": "Price type not specified in formula",
                "description": "Formula uses 'current_price' but doesn't specify if using close, VWAP, mid-price, etc.",
                "location": "Section 1, formula",
                "suggested_fix": "Specify: 'current_price = daily closing price at 16:00 ET'",
                "confidence": 0.85
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Mean calculation method not specified",
                "description": "Could be arithmetic mean, geometric mean, or other",
                "location": "Section 1, formula",
                "suggested_fix": "Specify: 'mean_price = arithmetic mean of daily closing prices over lookback period'",
                "confidence": 0.80
            }
        ],
        "self_containment_score": 0.40,
        "implementability_score": 0.35
    }
    
    save_document(doc, "doc_12_formula_precision", metadata)


# Document 13: Data specification gaps
def create_doc_13_data_gaps():
    """Document missing critical data specifications"""
    doc = Document()
    
    add_heading(doc, "Emerging Markets Local Currency Bond Strategy", 1)
    
    add_heading(doc, "1. Universe", 2)
    add_paragraph(doc, """
Invest in government bonds from emerging market countries, denominated in local currencies.
""")
    
    add_heading(doc, "2. Country Selection", 2)
    add_paragraph(doc, """
Include countries from the emerging markets list with minimum credit rating and sufficient market depth.
""")
    
    add_heading(doc, "3. Bond Selection", 2)
    add_paragraph(doc, """
Select bonds with maturity between 5 and 15 years.
Prices obtained from pricing service.
Yields calculated using standard methodology.
""")
    
    metadata = {
        "document_id": "doc_13_data_gaps",
        "title": "Emerging Markets Local Currency Bond Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "medium",
        "issue_count": 7,
        "issues": [
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Emerging markets list source not specified",
                "description": "'emerging markets list' referenced but source not identified (MSCI, FTSE, JP Morgan, etc.)",
                "location": "Section 2 - Country Selection",
                "suggested_fix": "Specify: 'Countries included in JP Morgan EMBI Global Diversified Index'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Minimum credit rating not specified",
                "description": "'minimum credit rating' mentioned but threshold not defined",
                "location": "Section 2 - Country Selection",
                "suggested_fix": "Specify: 'Minimum sovereign credit rating of B- (S&P) or B3 (Moody's)'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Market depth criteria not defined",
                "description": "'sufficient market depth' is subjective without quantitative threshold",
                "location": "Section 2 - Country Selection",
                "suggested_fix": "Specify: 'Minimum $5 billion outstanding local currency government debt and $100 million average daily trading volume'",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Pricing service not identified",
                "description": "'pricing service' mentioned but vendor not specified",
                "location": "Section 3 - Bond Selection",
                "suggested_fix": "Specify: 'Bond prices from Bloomberg Generic Prices (BGN) or Interactive Data Evaluated Pricing Service (IDEPS)'",
                "confidence": 0.90
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Yield calculation methodology not specified",
                "description": "'standard methodology' could mean yield-to-maturity, yield-to-call, or other",
                "location": "Section 3 - Bond Selection",
                "suggested_fix": "Specify: 'Yield-to-maturity calculated using semi-annual compounding convention or specify country-specific conventions'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "data_source_unspecified",
                "title": "Credit rating source not specified",
                "description": "References credit ratings but doesn't specify rating agency (S&P, Moody's, Fitch)",
                "location": "Section 2 - Country Selection",
                "suggested_fix": "Specify: 'Credit ratings from S&P Global Ratings; if unavailable, use Moody's or Fitch composite rating'",
                "confidence": 0.88
            },
            {
                "severity": "medium",
                "category": "data_source_unspecified",
                "title": "Trading volume data source not specified",
                "description": "Market depth assessment requires volume data but source not identified",
                "location": "Section 2 - Country Selection",
                "suggested_fix": "Specify: 'Trading volume data from Bloomberg or local exchange reporting'",
                "confidence": 0.82
            }
        ],
        "undefined_dependencies": [
            {"name": "Emerging markets classification", "type": "reference_data", "specs_missing": ["Provider", "Update frequency"]},
            {"name": "Bond pricing service", "type": "data_feed", "specs_missing": ["Vendor", "Pricing methodology", "Update time"]},
            {"name": "Credit ratings", "type": "external_system", "specs_missing": ["Rating agency", "Update frequency"]},
            {"name": "Trading volume data", "type": "data_feed", "specs_missing": ["Data source", "Aggregation method"]}
        ],
        "self_containment_score": 0.35,
        "implementability_score": 0.30
    }
    
    save_document(doc, "doc_13_data_gaps", metadata)


# Document 14: Complex multi-issue document
def create_doc_14_complex_multi():
    """Document with multiple severe issues across categories"""
    doc = Document()
    
    add_heading(doc, "Global Macro Tactical Asset Allocation Model", 1)
    
    add_heading(doc, "1. Asset Classes", 2)
    add_paragraph(doc, """
Allocate across equities, bonds, commodities, and alternatives based on macro signals.
""")
    
    add_heading(doc, "2. Signal Generation", 2)
    add_paragraph(doc, """
Signals derived from economic indicators, technical analysis, and sentiment measures.
Weight changes when signals reach significant levels.
""")
    
    add_heading(doc, "3. Implementation", 2)
    add_paragraph(doc, """
Rebalance periodically. Use derivatives where appropriate for efficiency.
Consult external manager for complex trades.
""")
    
    metadata = {
        "document_id": "doc_14_complex_multi",
        "title": "Global Macro Tactical Asset Allocation Model",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 15,
        "issues": [
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Asset class definitions missing",
                "description": "'equities, bonds, commodities, alternatives' mentioned but no specification of what instruments represent each",
                "location": "Section 1",
                "suggested_fix": "Define: Equities = MSCI ACWI, Bonds = Bloomberg Global Agg, Commodities = Bloomberg Commodity Index, Alternatives = specify exact instruments",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Economic indicators not specified",
                "description": "'economic indicators' referenced but no list of which indicators",
                "location": "Section 2",
                "suggested_fix": "List specific indicators: GDP growth, inflation (CPI), unemployment rate, PMI, yield curve, etc.",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "Technical analysis methodology not defined",
                "description": "'technical analysis' mentioned but no specification of indicators or rules",
                "location": "Section 2",
                "suggested_fix": "Specify: Moving average crossovers (50/200 day), RSI thresholds, or other specific technical indicators with parameters",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Sentiment measures not specified",
                "description": "'sentiment measures' referenced without identifying sources",
                "location": "Section 2",
                "suggested_fix": "Specify: AAII Sentiment Survey, VIX level, put/call ratios, or other specific sentiment indicators",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Significant signal levels not defined",
                "description": "'significant levels' is subjective without numerical thresholds",
                "location": "Section 2",
                "suggested_fix": "Define: 'Significant when composite signal exceeds ±1.5 standard deviations from historical mean'",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Signal combination methodology missing",
                "description": "Multiple signal types mentioned but no formula for combining them",
                "location": "Section 2",
                "suggested_fix": "Provide formula: Composite_Signal = w1×Economic + w2×Technical + w3×Sentiment, with weights specified",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "Weight change magnitude not specified",
                "description": "'Weight changes' but doesn't specify how much allocation shifts",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Allocation shifts by 10% of portfolio toward highest-signal asset class'",
                "confidence": 0.85
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Rebalancing frequency not specified",
                "description": "'periodically' is vague",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Rebalance monthly on first business day' or other specific frequency",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "'Where appropriate' criteria not defined",
                "description": "Derivatives usage conditional on 'appropriate' without defining criteria",
                "location": "Section 3",
                "suggested_fix": "Define: 'Use futures when transaction costs < 15 bps vs physical securities or for leverage efficiency'",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "external_dependency",
                "title": "External manager not identified",
                "description": "'external manager' referenced but not identified",
                "location": "Section 3",
                "suggested_fix": "Specify: Name external manager or define 'complex trades' that trigger external manager consultation",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "missing_governance",
                "title": "No approval process for manager consultation",
                "description": "External manager involvement mentioned but no governance around this decision",
                "location": "Section 3",
                "suggested_fix": "Add: 'External manager consultation requires CIO approval for trades exceeding $10M notional'",
                "confidence": 0.80
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Baseline asset allocation not specified",
                "description": "TAA model without specifying neutral/strategic allocation",
                "location": "Section 1",
                "suggested_fix": "Define neutral allocation: 'Strategic baseline: 40% equities, 35% bonds, 15% commodities, 10% alternatives'",
                "confidence": 0.88
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "Asset allocation bounds not specified",
                "description": "No min/max limits on asset class weights",
                "location": "Section 1",
                "suggested_fix": "Specify ranges: 'Equities 20-60%, Bonds 15-55%, Commodities 0-30%, Alternatives 0-20%'",
                "confidence": 0.90
            },
            {
                "severity": "medium",
                "category": "data_source_unspecified",
                "title": "Data frequency not specified",
                "description": "Economic and technical data update frequency not defined",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Economic indicators updated monthly, technical indicators daily, sentiment weekly'",
                "confidence": 0.75
            },
            {
                "severity": "medium",
                "category": "compliance_gap",
                "title": "No leverage or derivatives limits disclosed",
                "description": "Mentions derivatives but no disclosure of leverage limits or derivatives exposure caps",
                "location": "Section 3",
                "suggested_fix": "Add: 'Maximum gross exposure 150% of NAV. Derivatives notional limited to 100% of portfolio value'",
                "confidence": 0.82
            }
        ],
        "self_containment_score": 0.15,
        "implementability_score": 0.10,
        "notes": "Severely underspecified across all dimensions - requires comprehensive revision"
    }
    
    save_document(doc, "doc_14_complex_multi", metadata)


# Document 15: Simple with easy fixes
def create_doc_15_simple_fix():
    """Document with few, straightforward issues"""
    doc = Document()
    
    add_heading(doc, "U.S. Treasury Bond Index", 1)
    
    add_heading(doc, "1. Index Composition", 2)
    add_paragraph(doc, """
The index includes U.S. Treasury bonds with maturities between 1 and 10 years.
Bonds are weighted by market capitalization. Data sourced from U.S. Department of Treasury.
""")
    
    add_heading(doc, "2. Rebalancing", 2)
    add_paragraph(doc, """
Monthly rebalancing on the last business day of each month using closing prices.
New issues added at next rebalancing following auction.
""")
    
    add_heading(doc, "3. Pricing", 2)
    add_paragraph(doc, """
Bonds priced at closing prices from Federal Reserve Bank of New York at market close.
""")
    
    metadata = {
        "document_id": "doc_15_simple_fix",
        "title": "U.S. Treasury Bond Index",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "low",
        "issue_count": 3,
        "issues": [
            {
                "severity": "low",
                "category": "undefined_parameter",
                "title": "Maturity range endpoint ambiguity",
                "description": "'between 1 and 10 years' - unclear if inclusive or exclusive of endpoints",
                "location": "Section 1",
                "suggested_fix": "Clarify: 'Maturities greater than or equal to 1 year and less than or equal to 10 years' or 'Maturities ≥1 year and ≤10 years'",
                "confidence": 0.75
            },
            {
                "severity": "low",
                "category": "ambiguous_methodology",
                "title": "Market close time not specified",
                "description": "'market close' could be interpreted as Treasury market close or equity market close",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Market close defined as 15:00 ET (Treasury market close)' or '16:00 ET'",
                "confidence": 0.72
            },
            {
                "severity": "low",
                "category": "ambiguous_methodology",
                "title": "Treatment of newly auctioned bonds unclear during month",
                "description": "'added at next rebalancing' but handling between auction and rebalancing not specified",
                "location": "Section 2",
                "suggested_fix": "Clarify: 'Newly auctioned bonds not included in index until next monthly rebalancing. No intra-month additions.'",
                "confidence": 0.70
            }
        ],
        "self_containment_score": 0.92,
        "implementability_score": 0.90,
        "notes": "Minor clarifications needed - easily implementable with minimal changes"
    }
    
    save_document(doc, "doc_15_simple_fix", metadata)


# Document 16: Complex calendar issues
def create_doc_16_calendar_complex():
    """Document with sophisticated calendar and timezone issues"""
    doc = Document()
    
    add_heading(doc, "Asia-Pacific Multi-Market Strategy", 1)
    
    add_heading(doc, "1. Markets Covered", 2)
    add_paragraph(doc, """
Strategy trades equities across Tokyo, Hong Kong, Sydney, Singapore, and Mumbai markets.
""")
    
    add_heading(doc, "2. Trading Schedule", 2)
    add_paragraph(doc, """
Signals calculated at end of each trading day for each market.
Orders placed at market open the next day. Positions held until signals reverse.
""")
    
    add_heading(doc, "3. Rebalancing", 2)
    add_paragraph(doc, """
Monthly rebalancing on the last Friday of each month. All markets rebalanced simultaneously.
""")
    
    metadata = {
        "document_id": "doc_16_calendar_complex",
        "title": "Asia-Pacific Multi-Market Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 12,
        "issues": [
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Market-specific closing times not specified",
                "description": "'end of trading day' differs across markets - Tokyo 15:00 JST, Hong Kong 16:00 HKT, etc.",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Tokyo: 15:00 JST, Hong Kong: 16:00 HKT, Sydney: 16:00 AEDT/AEST, Singapore: 17:00 SGT, Mumbai: 15:30 IST'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Timezone coordination not addressed",
                "description": "Five markets in different timezones - signal timing relative to each market not clear",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Signals calculated using local market closing prices on same calendar date in each market's local timezone'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Market holiday handling not specified",
                "description": "Different markets have different holidays - behavior when some markets closed not defined",
                "location": "Section 2",
                "suggested_fix": "Add: 'On local market holidays, no trades executed for that market. Signals recalculated on next trading day using most recent available prices'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Simultaneous rebalancing impossible across timezones",
                "description": "'All markets rebalanced simultaneously' physically impossible given timezone differences",
                "location": "Section 3",
                "suggested_fix": "Clarify: 'Rebalancing executed at market open in each market on last Friday of month (local timezone), or next trading day if Friday is market holiday'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Last Friday definition ambiguous across timezones",
                "description": "Last Friday of month in Tokyo timezone differs from Hong Kong/Sydney due to date line",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Last Friday determined in UTC. Rebalancing occurs on that calendar date in each market's local timezone'",
                "confidence": 0.90
            },
            {
                "severity": "high",
                "category": "market_calendar",
                "title": "Order placement timing not specified",
                "description": "'market open the next day' - opening times differ by market",
                "location": "Section 2",
                "suggested_fix": "Specify opening times: 'Tokyo: 09:00 JST, Hong Kong: 09:30 HKT, Sydney: 10:00 AEDT/AEST, Singapore: 09:00 SGT, Mumbai: 09:15 IST'",
                "confidence": 0.88
            },
            {
                "severity": "high",
                "category": "market_calendar",
                "title": "Trading day definition not provided",
                "description": "'next trading day' behavior when Friday is holiday not specified",
                "location": "Section 3",
                "suggested_fix": "Add: 'If last Friday is market holiday, rebalancing occurs on next trading day for that market'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "market_calendar",
                "title": "Daylight saving time not addressed",
                "description": "Sydney observes DST (AEDT/AEST transition) affecting coordination with other markets",
                "location": "Section 2",
                "suggested_fix": "Add: 'Timezone adjustments account for DST transitions in relevant markets (Australia). Fixed UTC offsets: Tokyo +9, Hong Kong +8, Singapore +8, Mumbai +5:30'",
                "confidence": 0.82
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Signal coordination timing not defined",
                "description": "Markets close at different times - unclear if all signals wait for last market or calculated independently",
                "location": "Section 2",
                "suggested_fix": "Clarify: 'Signals calculated independently for each market immediately after local market close'",
                "confidence": 0.80
            },
            {
                "severity": "medium",
                "category": "market_calendar",
                "title": "Half-day trading sessions not addressed",
                "description": "Some markets have half-day trading before holidays - handling not specified",
                "location": "Section 2",
                "suggested_fix": "Add: 'On half-day trading sessions, signals calculated using half-day closing price at earlier local close time'",
                "confidence": 0.75
            },
            {
                "severity": "medium",
                "category": "market_calendar",
                "title": "Trading halt procedures not specified",
                "description": "Circuit breakers or trading halts would affect signal calculation",
                "location": "Section 2",
                "suggested_fix": "Add: 'If trading halted, use last available price. If market remains closed next day, no new orders for that market'",
                "confidence": 0.70
            },
            {
                "severity": "medium",
                "category": "data_source_unspecified",
                "title": "Calendar source not specified",
                "description": "Need authoritative source for market holiday calendars across five exchanges",
                "location": "Throughout",
                "suggested_fix": "Specify: 'Market calendars from respective exchange official sources: JPX (Tokyo), HKEX (Hong Kong), ASX (Sydney), SGX (Singapore), NSE (Mumbai)'",
                "confidence": 0.78
            }
        ],
        "self_containment_score": 0.25,
        "implementability_score": 0.20,
        "notes": "Critical ADR-018 example - complex timezone and multi-market calendar coordination issues"
    }
    
    save_document(doc, "doc_16_calendar_complex", metadata)


# Document 17: Moderate complexity
def create_doc_17_moderate():
    """Document with moderate issues, balanced complexity"""
    doc = Document()
    
    add_heading(doc, "U.S. Dividend Aristocrats Index", 1)
    
    add_heading(doc, "1. Eligibility", 2)
    add_paragraph(doc, """
S&P 500 companies that have increased dividends for at least 25 consecutive years.
""")
    
    add_heading(doc, "2. Weighting", 2)
    add_paragraph(doc, """
Constituents weighted by dividend yield with adjustments for company size.
Maximum single constituent weight capped.
""")
    
    add_heading(doc, "3. Maintenance", 2)
    add_paragraph(doc, """
Reviewed annually. Companies removed if they cut or suspend dividends.
Additions occur at the annual reconstitution.
""")
    
    metadata = {
        "document_id": "doc_17_moderate",
        "title": "U.S. Dividend Aristocrats Index",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "medium",
        "issue_count": 5,
        "issues": [
            {
                "severity": "medium",
                "category": "undefined_parameter",
                "title": "Dividend increase measurement not specified",
                "description": "'increased dividends' - unclear if absolute or percentage increase, and measurement basis",
                "location": "Section 1",
                "suggested_fix": "Specify: 'Increased annual dividend per share by any positive amount year-over-year for 25 consecutive fiscal years'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Size adjustment methodology not specified",
                "description": "'adjustments for company size' without defining the adjustment formula",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Weight = (Dividend_Yield × Market_Cap^0.5) / Sum of all (Dividend_Yield × Market_Cap^0.5)' or other explicit formula",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Maximum weight cap not specified",
                "description": "'capped' but cap level not defined",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Maximum 5% weight per constituent at reconstitution'",
                "confidence": 0.92
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "Annual reconstitution date not specified",
                "description": "'annually' but specific month/day not defined",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Annual reconstitution effective after close on third Friday in December'",
                "confidence": 0.90
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Dividend cut vs freeze distinction unclear",
                "description": "'cut or suspend' - unclear if maintaining flat dividend (not increasing) triggers removal",
                "location": "Section 3",
                "suggested_fix": "Clarify: 'Removal triggered by dividend reduction or suspension. Flat dividend (no increase) does not trigger removal, but company cannot qualify for future addition until 25-year streak resumes'",
                "confidence": 0.80
            }
        ],
        "self_containment_score": 0.65,
        "implementability_score": 0.60
    }
    
    save_document(doc, "doc_17_moderate", metadata)


# Document 18: Edge cases
def create_doc_18_edge_cases():
    """Document failing to specify edge case handling"""
    doc = Document()
    
    add_heading(doc, "Covered Call Writing Strategy", 1)
    
    add_heading(doc, "1. Strategy", 2)
    add_paragraph(doc, """
Buy underlying stocks and sell call options against positions.
Roll options at expiration if necessary.
""")
    
    add_heading(doc, "2. Strike Selection", 2)
    add_paragraph(doc, """
Sell calls at strike price above current stock price. Use standard option expiration dates.
""")
    
    add_heading(doc, "3. Position Management", 2)
    add_paragraph(doc, """
Close positions if called away. Replace with new positions.
""")
    
    metadata = {
        "document_id": "doc_18_edge_cases",
        "title": "Covered Call Writing Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 9,
        "issues": [
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Strike price offset not specified",
                "description": "'above current stock price' - how far above not defined",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Strike price 5% above current stock price, rounded to nearest standard strike'",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Option expiration period not specified",
                "description": "'standard expiration dates' - weekly, monthly, quarterly not specified",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Monthly options expiring on third Friday of each month, 30-45 days to expiration at initiation'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "'If necessary' roll criteria not defined",
                "description": "When to roll vs let assignment occur not specified",
                "location": "Section 1",
                "suggested_fix": "Define: 'Roll if option in-the-money by <$2 at 5 days before expiration. Otherwise allow assignment'",
                "confidence": 0.88
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Assignment handling not detailed",
                "description": "'called away' handling - timing of replacement not specified",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Upon assignment, replacement position established within 2 trading days. New call sold when replacement position established'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Early assignment handling not addressed",
                "description": "Options can be assigned early (dividend capture) - handling not specified",
                "location": "Section 3",
                "suggested_fix": "Add: 'If early assignment occurs, follow standard assignment procedure. Monitor for high early assignment risk before ex-dividend dates'",
                "confidence": 0.82
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "Option contract quantity not specified",
                "description": "How many calls to sell per 100 shares not defined (1:1, partial coverage, etc.)",
                "location": "Section 1",
                "suggested_fix": "Specify: 'Sell 1 call contract per 100 shares owned (100% coverage)'",
                "confidence": 0.90
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Dividend impact on roll decision not addressed",
                "description": "Dividends affect call pricing and assignment risk",
                "location": "Section 1",
                "suggested_fix": "Add: 'Evaluate early assignment risk when ex-dividend date falls before expiration. Roll to post-dividend expiration if assignment risk high'",
                "confidence": 0.75
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Odd lot handling not specified",
                "description": "Position sizes not divisible by 100 shares - handling not addressed",
                "location": "Section 1",
                "suggested_fix": "Specify: 'Round down to nearest 100 shares for option coverage. Odd lots held uncovered or excluded from strategy'",
                "confidence": 0.78
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Corporate action handling not specified",
                "description": "Stock splits, mergers, spinoffs affect options - adjustments not specified",
                "location": "Section 3",
                "suggested_fix": "Add: 'Corporate actions: Follow OCC adjustment memos for option contract adjustments. Reestablish standard contracts at next expiration'",
                "confidence": 0.80
            }
        ],
        "self_containment_score": 0.40,
        "implementability_score": 0.45,
        "notes": "Common edge cases not addressed - would encounter issues in real-world execution"
    }
    
    save_document(doc, "doc_18_edge_cases", metadata)


# Document 19: Almost clean
def create_doc_19_almost_clean():
    """Document with only minor issues"""
    doc = Document()
    
    add_heading(doc, "Bloomberg Barclays U.S. Aggregate Bond Index Replication", 1)
    
    add_heading(doc, "1. Objective", 2)
    add_paragraph(doc, """
Replicate the Bloomberg Barclays U.S. Aggregate Bond Index using sampling methodology.
Index data from Bloomberg Index Services Limited, updated daily at 17:00 ET.
""")
    
    add_heading(doc, "2. Sampling Method", 2)
    add_paragraph(doc, """
Select representative sample of 300-400 bonds from index constituents.
Maintain sector, duration, and quality weights within ±0.5% of benchmark.
Rebalance monthly using prices as of last business day of month.
""")
    
    add_heading(doc, "3. Performance Monitoring", 2)
    add_paragraph(doc, """
Track tracking error daily. Target tracking error below 15 basis points annualized.
Review bond selection quarterly to ensure continued representativeness.
""")
    
    metadata = {
        "document_id": "doc_19_almost_clean",
        "title": "Bloomberg Barclays U.S. Aggregate Bond Index Replication",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "low",
        "issue_count": 2,
        "issues": [
            {
                "severity": "low",
                "category": "undefined_parameter",
                "title": "Bond count target range could be more precise",
                "description": "'300-400 bonds' is a range - specific target or decision rule for exact count not provided",
                "location": "Section 2",
                "suggested_fix": "Consider specifying: 'Target 350 bonds, range 300-400 based on optimization results' or provide decision rule for selecting exact count",
                "confidence": 0.70
            },
            {
                "severity": "low",
                "category": "ambiguous_methodology",
                "title": "Tracking error measurement window not specified",
                "description": "'annualized' tracking error but rolling window period not defined",
                "location": "Section 3",
                "suggested_fix": "Specify: '15 basis points annualized calculated on 60-day rolling window' or other specific measurement period",
                "confidence": 0.72
            }
        ],
        "self_containment_score": 0.93,
        "implementability_score": 0.92,
        "notes": "Very minor issues - essentially production-ready with tiny clarifications"
    }
    
    save_document(doc, "doc_19_almost_clean", metadata)


# Document 20: Worst case scenario
def create_doc_20_worst_case():
    """Severely problematic document with numerous critical issues"""
    doc = Document()
    
    add_heading(doc, "Experimental Cryptocurrency Momentum Strategy", 1)
    
    add_heading(doc, "1. Universe", 2)
    add_paragraph(doc, """
Trade major cryptocurrencies based on momentum signals.
""")
    
    add_heading(doc, "2. Signals", 2)
    add_paragraph(doc, """
Buy tokens showing strong upward movement. Sell when momentum fades.
Use leverage to boost returns.
""")
    
    add_heading(doc, "3. Execution", 2)
    add_paragraph(doc, """
Trade on exchanges using algorithms. Monitor social media for sentiment.
Rebalance frequently.
""")
    
    metadata = {
        "document_id": "doc_20_worst_case",
        "title": "Experimental Cryptocurrency Momentum Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 20,
        "issues": [
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "'Major cryptocurrencies' not defined",
                "description": "No specification of which tokens qualify",
                "location": "Section 1",
                "suggested_fix": "Define universe: 'Top 20 cryptocurrencies by market capitalization from CoinMarketCap, excluding stablecoins'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Market cap data source not specified",
                "description": "Need source for determining 'major' status",
                "location": "Section 1",
                "suggested_fix": "Specify: 'Market capitalization data from CoinMarketCap or CoinGecko API, updated hourly'",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "'Strong upward movement' not quantified",
                "description": "Completely subjective without numerical threshold",
                "location": "Section 2",
                "suggested_fix": "Define: 'Strong movement = 20% gain over past 7 days AND positive 24-hour momentum'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "'Momentum fades' not defined",
                "description": "Exit criteria completely unspecified",
                "location": "Section 2",
                "suggested_fix": "Define: 'Momentum fades when 7-day return drops below 5% OR enters negative territory'",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Leverage amount not specified",
                "description": "'use leverage' without specifying magnitude",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Maximum 2x leverage using perpetual futures. Gross exposure capped at 200% of NAV'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "external_dependency",
                "title": "Exchanges not identified",
                "description": "'exchanges' mentioned but none specified",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Primary: Coinbase Pro, Binance, Kraken. Execution split across exchanges for liquidity'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Pricing data source not specified",
                "description": "Need authoritative price source for signals and execution",
                "location": "Throughout",
                "suggested_fix": "Specify: 'Pricing from CryptoCompare CCCAGG (aggregate) or exchange-specific APIs with median calculation'",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "Algorithm execution strategy not specified",
                "description": "'using algorithms' without defining execution strategy",
                "location": "Section 3",
                "suggested_fix": "Define: 'TWAP execution over 15-minute window. Limit orders at midpoint ± 10 bps. Market orders if not filled within 30 minutes'",
                "confidence": 0.88
            },
            {
                "severity": "critical",
                "category": "data_source_unspecified",
                "title": "Social media sources not identified",
                "description": "'social media' monitoring without specifying platforms or metrics",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Twitter sentiment analysis using LunarCrush API or specify manual vs automated monitoring'",
                "confidence": 0.85
            },
            {
                "severity": "critical",
                "category": "ambiguous_methodology",
                "title": "Sentiment impact on strategy not defined",
                "description": "Mentions monitoring sentiment but no specification of how it affects decisions",
                "location": "Section 3",
                "suggested_fix": "Define: 'Negative sentiment override: Exit position if sentiment score drops below -0.5 on scale of -1 to +1'",
                "confidence": 0.82
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "'Frequently' rebalance not quantified",
                "description": "Rebalancing frequency completely unspecified",
                "location": "Section 3",
                "suggested_fix": "Specify: 'Rebalance daily at 00:00 UTC' or 'Intraday rebalancing every 6 hours'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "compliance_gap",
                "title": "No risk disclosures for crypto volatility",
                "description": "Crypto strategies require extensive risk warnings - completely absent",
                "location": "Throughout",
                "suggested_fix": "Add comprehensive risk section: Extreme volatility (50%+ daily moves), total loss potential, regulatory uncertainty, exchange counterparty risk, liquidity risk",
                "confidence": 0.98
            },
            {
                "severity": "critical",
                "category": "compliance_gap",
                "title": "No leverage risk disclosure",
                "description": "Leveraged crypto extremely risky - no warnings",
                "location": "Section 2",
                "suggested_fix": "Add: 'Leverage Risk: 2x leverage amplifies losses. Liquidation possible with 50% adverse move. Margin calls may force liquidation at worst prices during volatile periods'",
                "confidence": 0.97
            },
            {
                "severity": "critical",
                "category": "compliance_gap",
                "title": "No custody or counterparty risk disclosure",
                "description": "Exchange risk is critical for crypto - no mention",
                "location": "Section 3",
                "suggested_fix": "Add: 'Counterparty Risk: Assets held on exchanges subject to exchange insolvency, hacking, or operational failures. Not FDIC or SIPC insured'",
                "confidence": 0.95
            },
            {
                "severity": "high",
                "category": "missing_governance",
                "title": "No position size limits",
                "description": "No specification of position sizing or concentration limits",
                "location": "Throughout",
                "suggested_fix": "Add: 'Maximum 20% allocation per token. Minimum liquidity requirement: $100M 24-hour volume'",
                "confidence": 0.90
            },
            {
                "severity": "high",
                "category": "missing_governance",
                "title": "No stop loss or drawdown controls",
                "description": "No risk management parameters specified",
                "location": "Throughout",
                "suggested_fix": "Add: 'Individual position stop loss: 25%. Portfolio stop loss: 40% drawdown triggers deleveraging. Maximum daily loss: 10%'",
                "confidence": 0.92
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "No specification of momentum calculation period",
                "description": "Momentum mentioned but lookback period not defined",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Momentum calculated using 7-day, 14-day, and 30-day returns. Composite signal = weighted average'",
                "confidence": 0.88
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Return calculation method not specified",
                "description": "Crypto trades 24/7 - daily return calculation time not defined",
                "location": "Section 2",
                "suggested_fix": "Specify: 'Returns calculated using prices at 00:00 UTC daily. Intraday volatility monitored separately'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "compliance_gap",
                "title": "No regulatory status disclosure",
                "description": "Many crypto assets have unclear regulatory status",
                "location": "Throughout",
                "suggested_fix": "Add: 'Regulatory Risk: Cryptocurrencies subject to evolving regulation. Some tokens may be classified as securities. Strategy may need to terminate if regulations change'",
                "confidence": 0.88
            },
            {
                "severity": "medium",
                "category": "data_source_unspecified",
                "title": "No specification of data reliability or validation",
                "description": "Crypto data can be unreliable - no data quality measures",
                "location": "Throughout",
                "suggested_fix": "Add: 'Data Quality: Use median of 3+ exchange prices. Discard outliers >10% from median. Halt trading if price feeds unavailable from 2+ exchanges'",
                "confidence": 0.82
            }
        ],
        "compliance_assessment": {
            "sec_compliance": "non_compliant",
            "investor_suitability": "unsuitable_without_major_revisions",
            "reason": "Missing critical risk disclosures, undefined strategy parameters, regulatory concerns",
            "required_changes": "Comprehensive rewrite required"
        },
        "self_containment_score": 0.10,
        "implementability_score": 0.05,
        "notes": "SEVERE ISSUES - Nearly impossible to implement without comprehensive specification. High risk of regulatory non-compliance. Unsuitable for investor presentation in current form."
    }
    
    save_document(doc, "doc_20_worst_case", metadata)


def main():
    print("Generating final test documents (Part 3)...")
    print(f"Output directory: {OUTPUT_DIR}\n")
    
    create_doc_11_perfect_compliance()
    create_doc_12_formula_precision()
    create_doc_13_data_gaps()
    create_doc_14_complex_multi()
    create_doc_15_simple_fix()
    create_doc_16_calendar_complex()
    create_doc_17_moderate()
    create_doc_18_edge_cases()
    create_doc_19_almost_clean()
    create_doc_20_worst_case()
    
    print(f"\n✓ Successfully created documents 11-20 in {OUTPUT_DIR}")
    print(f"\n📊 Complete Test Suite Summary:")
    print(f"   Documents 01-05: Baseline + primary issue categories")
    print(f"   Documents 06-10: External dependencies, governance, consistency, ambiguity, compliance")
    print(f"   Documents 11-20: Full spectrum from perfect to severely flawed")
    print(f"\n   Total: 20 .docx documents + 20 .json metadata files = 40 files")


if __name__ == "__main__":
    main()
