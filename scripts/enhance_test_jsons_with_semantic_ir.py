#!/usr/bin/env python3
"""
Script to enhance test document JSON files with expected semantic IR data.

This script reads each .docx file in the test documents directory, extracts
semantic content (formulas, definitions, tables, cross-references), and adds
this as expected_semantic_ir to the corresponding JSON file.
"""
import json
import re
from pathlib import Path
from typing import List, Dict, Any
from docx import Document


TEST_DOCS_DIR = Path("/workspaces/data/test_documents")


def extract_formulas(doc: Document) -> List[Dict[str, Any]]:
    """Extract mathematical formulas from document."""
    formulas = []
    formula_id = 1

    # Patterns to identify formulas
    formula_indicators = ['=', '×', '÷', '/', '+', '-', '∑', '∏', '√']

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Skip empty paragraphs or headers
        if not text or para.style.name.startswith('Heading'):
            continue

        # Check if paragraph contains formula indicators
        if any(indicator in text for indicator in formula_indicators):
            # Try to extract formula expression
            if '=' in text:
                # Split on = to get formula components
                parts = text.split('=', 1)
                if len(parts) == 2:
                    formula_name = parts[0].strip()
                    formula_expr = parts[1].strip()

                    # Extract variables (single uppercase letters or camelCase identifiers)
                    var_pattern = r'\b[A-Z][a-z]*[A-Z]?[a-z]*\b|[w]\d+'
                    variables = list(set(re.findall(var_pattern, formula_expr)))

                    formulas.append({
                        "id": f"formula_{formula_id}",
                        "name": formula_name if len(formula_name) < 50 else None,
                        "expression": text,
                        "location": f"Paragraph {para_idx + 1}",
                        "variables": [{"name": v, "source": None} for v in variables if v not in ['Sum', 'Max', 'Min']],
                        "parameters": []
                    })
                    formula_id += 1

    return formulas


def extract_definitions(doc: Document) -> List[Dict[str, Any]]:
    """Extract term definitions from document."""
    definitions = []
    def_id = 1

    # Definition patterns
    def_patterns = [
        r'"([^"]+)"\s+means\s+(.+)',
        r'"([^"]+)"\s+(?:refers to|is defined as)\s+(.+)',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+is\s+(.+)',
        r'Where\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+is\s+(.+)',
    ]

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text or para.style.name.startswith('Heading'):
            continue

        # Try each pattern
        for pattern in def_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                term = match.group(1).strip()
                definition = match.group(2).strip()

                # Clean up definition (remove trailing punctuation from regex capture)
                definition = re.sub(r'[.;,]$', '', definition)

                definitions.append({
                    "id": f"term_{def_id}",
                    "term": term,
                    "definition": definition,
                    "location": f"Paragraph {para_idx + 1}",
                    "aliases": []
                })
                def_id += 1
                break  # Only match one pattern per paragraph

    return definitions


def extract_tables(doc: Document) -> List[Dict[str, Any]]:
    """Extract tables from document."""
    tables = []

    for table_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # Extract headers from first row
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Extract data rows
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)

        # Identify parameter tables (usually have "Parameter", "Value", "Description" columns)
        is_parameter_table = any(
            h.lower() in ['parameter', 'value', 'description', 'variable']
            for h in headers
        )

        table_data = {
            "id": f"table_{table_idx + 1}",
            "title": f"Table {table_idx + 1}",
            "headers": headers,
            "row_count": len(rows),
            "column_count": len(headers),
            "location": f"Table {table_idx + 1}",
            "is_parameter_table": is_parameter_table
        }

        # If it's a parameter table, extract parameters
        if is_parameter_table and len(rows) > 0:
            table_data["provides_parameters"] = []
            # Try to find parameter/value columns
            param_col = next((i for i, h in enumerate(headers) if 'param' in h.lower()), 0)
            value_col = next((i for i, h in enumerate(headers) if 'value' in h.lower()), 1 if len(headers) > 1 else None)

            if value_col is not None:
                for row in rows[:5]:  # Sample first 5 parameters
                    if len(row) > max(param_col, value_col):
                        table_data["provides_parameters"].append({
                            "name": row[param_col],
                            "value": row[value_col] if value_col < len(row) else None
                        })

        tables.append(table_data)

    return tables


def extract_cross_references(doc: Document) -> List[Dict[str, Any]]:
    """Extract cross-references from document."""
    cross_refs = []
    ref_id = 1

    # Cross-reference patterns
    ref_patterns = [
        r'(?:see|refer to|as (?:described|defined) in)\s+(?:Section|Appendix|Table|Figure)\s+(\d+|[A-Z])',
        r'(?:Section|Appendix|Table|Figure)\s+(\d+|[A-Z])',
    ]

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        for pattern in ref_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                target = match.group(0).strip()
                cross_refs.append({
                    "id": f"ref_{ref_id}",
                    "source_location": f"Paragraph {para_idx + 1}",
                    "target": target,
                    "reference_text": match.group(0)
                })
                ref_id += 1

    return cross_refs


def analyze_document(docx_path: Path) -> Dict[str, Any]:
    """Analyze a document and extract semantic IR."""
    doc = Document(docx_path)

    return {
        "definitions": extract_definitions(doc),
        "formulae": extract_formulas(doc),
        "tables": extract_tables(doc),
        "cross_references": extract_cross_references(doc)
    }


def enhance_json_files():
    """Enhance all JSON files with expected semantic IR."""
    docx_files = sorted(TEST_DOCS_DIR.glob("*.docx"))

    for docx_file in docx_files:
        json_file = docx_file.with_suffix(".json")

        if not json_file.exists():
            print(f"⚠️  No JSON file for {docx_file.name}")
            continue

        print(f"\n📄 Processing {docx_file.name}...")

        # Load existing JSON
        with open(json_file, 'r') as f:
            data = json.load(f)

        # Extract semantic IR
        semantic_ir = analyze_document(docx_file)

        # Add to JSON
        data["expected_semantic_ir"] = semantic_ir

        # Print summary
        print(f"   ✓ Definitions: {len(semantic_ir['definitions'])}")
        print(f"   ✓ Formulae: {len(semantic_ir['formulae'])}")
        print(f"   ✓ Tables: {len(semantic_ir['tables'])}")
        print(f"   ✓ Cross-references: {len(semantic_ir['cross_references'])}")

        # Save enhanced JSON
        with open(json_file, 'w') as f:
            json.dump(data, f, indent=2)

        print(f"   ✅ Updated {json_file.name}")


def main():
    """Main entry point."""
    print("=" * 70)
    print("Enhancing Test Document JSON Files with Semantic IR")
    print("=" * 70)

    enhance_json_files()

    print("\n" + "=" * 70)
    print("✅ Complete!")
    print("=" * 70)


if __name__ == "__main__":
    main()
