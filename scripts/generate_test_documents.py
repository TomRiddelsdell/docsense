#!/usr/bin/env python3
"""
Generate test documents for demonstration and testing purposes.

Creates 20 .docx files with varying complexity and issues, plus metadata JSON files
describing what issues should be detected.
"""

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Output directory
OUTPUT_DIR = Path("data/test_documents")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, style='Normal'):
    """Add a paragraph to the document"""
    para = doc.add_paragraph(text, style=style)
    return para


def add_formula(doc, formula_text):
    """Add a formula (as text for now)"""
    para = doc.add_paragraph()
    run = para.add_run(formula_text)
    run.italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return para


def save_document(doc, filename, metadata):
    """Save document and metadata"""
    doc_path = OUTPUT_DIR / f"{filename}.docx"
    json_path = OUTPUT_DIR / f"{filename}.json"
    
    doc.save(doc_path)
    
    with open(json_path, 'w') as f:
        json.dump(metadata, f, indent=2)
    
    print(f"Created: {doc_path.name} and {json_path.name}")


# Document 1: Clean Document (Minimal Issues)
def create_doc_01_clean():
    """Nearly perfect document with minimal issues - baseline for testing"""
    doc = Document()
    
    add_heading(doc, "S&P Technology Sector Equal Weight Index Methodology", 1)
    
    add_heading(doc, "1. Index Overview", 2)
    add_paragraph(doc, """
The S&P Technology Sector Equal Weight Index provides exposure to U.S. technology companies 
with equal weighting applied to all constituents. The index is rebalanced quarterly on the 
third Friday of March, June, September, and December.
""")
    
    add_heading(doc, "2. Constituent Selection", 2)
    add_paragraph(doc, """
Constituents are selected from the S&P 500 Index and must meet the following criteria:
- Market capitalization greater than $5 billion USD
- Average daily trading volume greater than $10 million USD over the past 3 months
- Classification in GICS Technology Sector (Code: 45)
- Minimum 12 months of trading history
""")
    
    add_heading(doc, "3. Weighting Methodology", 2)
    add_paragraph(doc, """
All constituents receive equal weight at each rebalancing date. The equal weight is calculated as:

Weight per constituent = 100% / Number of constituents

As of the December 2024 rebalancing, there are 75 constituents, resulting in each constituent 
having a weight of 1.333%.
""")
    
    add_heading(doc, "4. Rebalancing Procedure", 2)
    add_paragraph(doc, """
The index is rebalanced quarterly according to the following schedule:
- Effective Date: Third Friday of March, June, September, December
- Reference Date: Close of business 5 NYSE trading days prior to effective date
- Announcement Date: 10 NYSE trading days prior to effective date

On the effective date at market close (16:00 America/New_York):
1. Calculate current constituent weights
2. Determine target equal weights (1/N for each constituent)
3. Calculate required trades to reach target weights
4. Execute rebalancing at closing prices
""")
    
    add_heading(doc, "5. Corporate Actions", 2)
    add_paragraph(doc, """
Corporate actions are handled as follows:

Dividends: Dividends are reinvested on the ex-dividend date at the opening price.

Stock Splits: Shares outstanding are adjusted on the ex-date. No weight adjustment required 
as market capitalization remains constant.

Mergers and Acquisitions: 
- If acquirer is an index constituent: Combined entity weight equals sum of both constituents
- If acquirer is not a constituent: Remove acquired company, redistribute weight pro-rata to remaining constituents
- Effective on completion date as announced by the exchange

Spin-offs: Parent company weight is split between parent and spin-off based on market 
capitalizations on the first trading day of the spin-off.
""")
    
    add_heading(doc, "6. Index Calculation", 2)
    add_paragraph(doc, """
The index level is calculated using the following formula:
""")
    add_formula(doc, "Index Level = (Current Market Value / Base Market Value) × Base Level")
    
    add_paragraph(doc, """
Where:
- Current Market Value = Sum of (Price_i × Shares_i × Weight_i) for all constituents
- Base Market Value = Market value at base date (January 1, 2020)
- Base Level = 1000 (set at inception)
- Price_i = Current price of constituent i in USD
- Shares_i = Outstanding shares of constituent i
- Weight_i = Current weight of constituent i
""")
    
    add_heading(doc, "7. Data Sources", 2)
    add_paragraph(doc, """
- Price Data: NYSE TAQ (Trade and Quote) database, end-of-day closing prices
- Market Capitalization: S&P Capital IQ, updated daily
- Corporate Actions: Bloomberg Corporate Actions Feed, verified against exchange announcements
- Trading Volume: NYSE TAQ, calculated as 90-day average daily dollar volume
- GICS Classifications: MSCI/S&P Global Industry Classification Standard, updated quarterly
""")
    
    metadata = {
        "document_id": "doc_01_clean",
        "title": "S&P Technology Sector Equal Weight Index Methodology",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "low",
        "issue_count": 1,
        "issues": [
            {
                "severity": "low",
                "category": "data_source_unspecified",
                "title": "NYSE trading day specification could be more explicit",
                "description": "Document mentions '5 NYSE trading days' but doesn't explicitly state this excludes NYSE holidays",
                "location": "Section 4 - Rebalancing Procedure",
                "original_text": "5 NYSE trading days prior to effective date",
                "suggested_fix": "5 NYSE trading days (excluding weekends and NYSE market holidays) prior to effective date",
                "confidence": 0.65
            }
        ],
        "strengths": [
            "Clear constituent selection criteria with specific thresholds",
            "Complete rebalancing procedure with specific dates and times",
            "Comprehensive corporate action handling",
            "Explicit index calculation formula with all variables defined",
            "All data sources specified with vendor names",
            "Timezone specified for market close",
            "All parameters have specific numeric values"
        ],
        "self_containment_score": 0.95,
        "implementability_score": 0.98
    }
    
    save_document(doc, "doc_01_clean", metadata)


# Document 2: Missing Appendix Reference
def create_doc_02_missing_appendix():
    """Document referencing external appendix that's not included"""
    doc = Document()
    
    add_heading(doc, "Emerging Markets Dividend Growth Strategy", 1)
    
    add_heading(doc, "1. Strategy Overview", 2)
    add_paragraph(doc, """
This strategy invests in emerging market equities with consistent dividend growth. 
The portfolio is constructed using a proprietary scoring model detailed in Appendix A.
""")
    
    add_heading(doc, "2. Universe Definition", 2)
    add_paragraph(doc, """
The investment universe consists of companies that meet the following criteria:
- Listed on exchanges in emerging market countries (see Appendix B for country list)
- Market capitalization above $500 million
- Minimum 5-year dividend payment history
- Dividend growth rate greater than inflation rate
""")
    
    add_heading(doc, "3. Portfolio Construction", 2)
    add_paragraph(doc, """
Stocks are scored based on:
- Dividend yield (40% weight)
- Dividend growth rate (30% weight)
- Payout ratio sustainability (30% weight)

The top 50 stocks by composite score are selected for the portfolio. 
Detailed scoring methodology is provided in Appendix C.
""")
    
    metadata = {
        "document_id": "doc_02_missing_appendix",
        "title": "Emerging Markets Dividend Growth Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "medium",
        "issue_count": 4,
        "issues": [
            {
                "severity": "critical",
                "category": "missing_reference",
                "title": "Appendix A referenced but not included",
                "description": "Document references 'Appendix A' for proprietary scoring model, which is essential for implementation but not included",
                "location": "Section 1 - Strategy Overview",
                "original_text": "proprietary scoring model detailed in Appendix A",
                "suggested_fix": "Include Appendix A or incorporate the scoring model details directly in the methodology section",
                "confidence": 0.98
            },
            {
                "severity": "critical",
                "category": "missing_reference",
                "title": "Appendix B referenced but not included",
                "description": "Document references 'Appendix B for country list' but appendix is not included. Cannot determine which countries qualify as emerging markets",
                "location": "Section 2 - Universe Definition",
                "original_text": "see Appendix B for country list",
                "suggested_fix": "List emerging market countries directly or include Appendix B",
                "confidence": 0.98
            },
            {
                "severity": "critical",
                "category": "missing_reference",
                "title": "Appendix C referenced but not included",
                "description": "Detailed scoring methodology in Appendix C is not included, making implementation impossible",
                "location": "Section 3 - Portfolio Construction",
                "original_text": "Detailed scoring methodology is provided in Appendix C",
                "suggested_fix": "Include Appendix C with complete scoring formulas and examples",
                "confidence": 0.98
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "Inflation rate source and calculation not specified",
                "description": "Document requires 'dividend growth rate greater than inflation rate' but doesn't specify which inflation measure or data source",
                "location": "Section 2 - Universe Definition",
                "original_text": "Dividend growth rate greater than inflation rate",
                "suggested_fix": "Specify: 'Dividend growth rate greater than the country-specific Consumer Price Index (CPI) as published by the IMF World Economic Outlook database'",
                "confidence": 0.90
            }
        ],
        "missing_documents": [
            {
                "referenced_name": "Appendix A - Proprietary Scoring Model",
                "reference_location": "Section 1",
                "purpose": "Details the proprietary scoring model",
                "criticality": "critical"
            },
            {
                "referenced_name": "Appendix B - Emerging Market Country List",
                "reference_location": "Section 2",
                "purpose": "Defines which countries qualify as emerging markets",
                "criticality": "critical"
            },
            {
                "referenced_name": "Appendix C - Detailed Scoring Methodology",
                "reference_location": "Section 3",
                "purpose": "Provides detailed scoring formulas and calculations",
                "criticality": "critical"
            }
        ],
        "self_containment_score": 0.20,
        "implementability_score": 0.15
    }
    
    save_document(doc, "doc_02_missing_appendix", metadata)


# Document 3: Undefined Parameters
def create_doc_03_undefined_parameters():
    """Document with multiple undefined parameters"""
    doc = Document()
    
    add_heading(doc, "Low Volatility Equity Portfolio Strategy", 1)
    
    add_heading(doc, "1. Objective", 2)
    add_paragraph(doc, """
Construct a portfolio of low-volatility equities that outperforms the benchmark 
with significantly lower risk.
""")
    
    add_heading(doc, "2. Stock Selection", 2)
    add_paragraph(doc, """
Select stocks from the investment universe where:
- Volatility is below the threshold
- Beta is less than the maximum allowed
- Liquidity exceeds the minimum requirement
- Market cap is above the cutoff
""")
    
    add_heading(doc, "3. Position Sizing", 2)
    add_paragraph(doc, """
Each position is weighted based on its inverse volatility. Maximum position size 
is capped to prevent concentration. Minimum position size is enforced to maintain 
meaningful exposure.
""")
    
    add_heading(doc, "4. Rebalancing", 2)
    add_paragraph(doc, """
The portfolio is rebalanced when drift exceeds the tolerance threshold or at 
regular intervals, whichever comes first. Transaction costs are considered, 
and rebalancing only occurs if expected benefit exceeds the cost.
""")
    
    metadata = {
        "document_id": "doc_03_undefined_parameters",
        "title": "Low Volatility Equity Portfolio Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "medium",
        "issue_count": 12,
        "issues": [
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Volatility threshold not specified",
                "description": "Requires 'volatility below the threshold' but threshold value is not defined",
                "location": "Section 2 - Stock Selection",
                "original_text": "Volatility is below the threshold",
                "suggested_fix": "Specify: 'Annualized volatility (60-day rolling standard deviation of daily returns) is below 15%'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Maximum beta not specified",
                "description": "Requires 'beta less than the maximum allowed' but maximum is not defined",
                "location": "Section 2 - Stock Selection",
                "original_text": "Beta is less than the maximum allowed",
                "suggested_fix": "Specify: 'Beta relative to S&P 500 is less than 0.8 (calculated using 12-month daily returns)'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Liquidity minimum requirement not specified",
                "description": "Requires 'liquidity exceeds the minimum requirement' but minimum is not defined",
                "location": "Section 2 - Stock Selection",
                "original_text": "Liquidity exceeds the minimum requirement",
                "suggested_fix": "Specify: 'Average daily trading volume exceeds $5 million over the past 90 calendar days'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Market cap cutoff not specified",
                "description": "Requires 'market cap above the cutoff' but cutoff value is not defined",
                "location": "Section 2 - Stock Selection",
                "original_text": "Market cap is above the cutoff",
                "suggested_fix": "Specify: 'Market capitalization exceeds $2 billion USD'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Maximum position size cap not specified",
                "description": "States 'maximum position size is capped' but doesn't specify the cap",
                "location": "Section 3 - Position Sizing",
                "original_text": "Maximum position size is capped to prevent concentration",
                "suggested_fix": "Specify: 'Maximum position size is capped at 5% of portfolio value'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Minimum position size not specified",
                "description": "States 'minimum position size is enforced' but doesn't specify the minimum",
                "location": "Section 3 - Position Sizing",
                "original_text": "Minimum position size is enforced",
                "suggested_fix": "Specify: 'Minimum position size is 0.5% of portfolio value'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "undefined_parameter",
                "title": "Drift tolerance threshold not specified",
                "description": "Rebalancing triggered when 'drift exceeds tolerance threshold' but threshold not defined",
                "location": "Section 4 - Rebalancing",
                "original_text": "drift exceeds the tolerance threshold",
                "suggested_fix": "Specify: 'drift exceeds 2 percentage points from target weight'",
                "confidence": 0.95
            },
            {
                "severity": "high",
                "category": "undefined_parameter",
                "title": "Regular rebalancing interval not specified",
                "description": "States 'regular intervals' but doesn't specify the interval",
                "location": "Section 4 - Rebalancing",
                "original_text": "regular intervals",
                "suggested_fix": "Specify: 'quarterly intervals (last trading day of March, June, September, December)'",
                "confidence": 0.90
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Benchmark not specified",
                "description": "Mentions 'outperforms the benchmark' but doesn't identify which benchmark",
                "location": "Section 1 - Objective",
                "original_text": "outperforms the benchmark",
                "suggested_fix": "Specify: 'outperforms the S&P 500 Total Return Index'",
                "confidence": 0.92
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Volatility calculation methodology not specified",
                "description": "References volatility but doesn't specify calculation method, lookback period, or frequency",
                "location": "Section 2 - Stock Selection",
                "original_text": "Volatility is below the threshold",
                "suggested_fix": "Specify: 'Annualized volatility calculated as standard deviation of daily returns over rolling 60-day window, using close-to-close returns'",
                "confidence": 0.88
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Transaction cost calculation not specified",
                "description": "Mentions 'transaction costs are considered' but doesn't specify how they're calculated",
                "location": "Section 4 - Rebalancing",
                "original_text": "Transaction costs are considered",
                "suggested_fix": "Specify: 'Transaction costs estimated as 0.1% of trade value (including commissions, market impact, and spread)'",
                "confidence": 0.85
            },
            {
                "severity": "medium",
                "category": "ambiguous_methodology",
                "title": "Cost-benefit calculation not specified",
                "description": "States 'expected benefit exceeds the cost' but doesn't define how benefit is calculated",
                "location": "Section 4 - Rebalancing",
                "original_text": "expected benefit exceeds the cost",
                "suggested_fix": "Specify: 'Expected tracking error reduction must exceed transaction costs by at least 0.05%'",
                "confidence": 0.80
            }
        ],
        "self_containment_score": 0.30,
        "implementability_score": 0.25
    }
    
    save_document(doc, "doc_03_undefined_parameters", metadata)


# Document 4: Incomplete Formulas
def create_doc_04_incomplete_formulas():
    """Document with incomplete or ambiguous formulas"""
    doc = Document()
    
    add_heading(doc, "Momentum Factor Index Methodology", 1)
    
    add_heading(doc, "1. Momentum Score Calculation", 2)
    add_paragraph(doc, """
The momentum score for each stock is calculated as:
""")
    add_formula(doc, "Momentum Score = (Current Price / Historical Price) - 1")
    
    add_paragraph(doc, """
Where Historical Price is the price from the lookback period.
""")
    
    add_heading(doc, "2. Risk-Adjusted Momentum", 2)
    add_paragraph(doc, """
Risk-adjusted momentum is calculated by dividing momentum by volatility:
""")
    add_formula(doc, "Risk-Adjusted Momentum = Momentum / Volatility")
    
    add_heading(doc, "3. Composite Score", 2)
    add_paragraph(doc, """
The final composite score combines multiple factors:
""")
    add_formula(doc, "Composite Score = w1 × Factor1 + w2 × Factor2 + w3 × Factor3")
    
    add_paragraph(doc, """
Where weights sum to 100%.
""")
    
    add_heading(doc, "4. Index Weight Calculation", 2)
    add_paragraph(doc, """
Each constituent's weight is proportional to its score:
""")
    add_formula(doc, "Weight = Score / Sum(Scores)")
    
    metadata = {
        "document_id": "doc_04_incomplete_formulas",
        "title": "Momentum Factor Index Methodology",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "medium",
        "issue_count": 10,
        "issues": [
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Lookback period for Historical Price not specified",
                "description": "Formula uses 'Historical Price from the lookback period' but lookback period duration is not specified",
                "location": "Section 1 - Momentum Score Calculation",
                "original_text": "Historical Price is the price from the lookback period",
                "suggested_fix": "Specify: 'Historical Price is the closing price from 12 months prior (252 NYSE trading days)'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Price definition ambiguous - which price?",
                "description": "Formula uses 'Current Price' and 'Historical Price' but doesn't specify if these are closing, opening, or average prices",
                "location": "Section 1 - Momentum Score Calculation",
                "original_text": "Current Price / Historical Price",
                "suggested_fix": "Specify: 'Closing prices adjusted for splits and dividends'",
                "confidence": 0.90
            },
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Volatility calculation not specified",
                "description": "Risk-adjusted momentum formula uses Volatility but doesn't specify how it's calculated",
                "location": "Section 2 - Risk-Adjusted Momentum",
                "original_text": "Momentum / Volatility",
                "suggested_fix": "Specify: 'Volatility is the annualized standard deviation of daily returns over the past 60 trading days'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Factor1, Factor2, Factor3 not defined",
                "description": "Composite score formula references Factor1, Factor2, Factor3 but these are not defined anywhere",
                "location": "Section 3 - Composite Score",
                "original_text": "w1 × Factor1 + w2 × Factor2 + w3 × Factor3",
                "suggested_fix": "Define each factor explicitly, e.g., 'Factor1 = Momentum Score, Factor2 = Risk-Adjusted Momentum, Factor3 = Value Score'",
                "confidence": 0.98
            },
            {
                "severity": "critical",
                "category": "incomplete_formula",
                "title": "Weights w1, w2, w3 not specified",
                "description": "Formula uses weights w1, w2, w3 but their values are not provided",
                "location": "Section 3 - Composite Score",
                "original_text": "w1 × Factor1 + w2 × Factor2 + w3 × Factor3",
                "suggested_fix": "Specify: 'w1 = 0.50, w2 = 0.30, w3 = 0.20'",
                "confidence": 0.98
            },
            {
                "severity": "high",
                "category": "incomplete_formula",
                "title": "Missing division by zero handling",
                "description": "Risk-adjusted momentum divides by volatility with no specification for zero volatility case",
                "location": "Section 2 - Risk-Adjusted Momentum",
                "original_text": "Momentum / Volatility",
                "suggested_fix": "Add: 'If Volatility = 0, set Risk-Adjusted Momentum = 0 or exclude from index'",
                "confidence": 0.85
            },
            {
                "severity": "high",
                "category": "incomplete_formula",
                "title": "Negative momentum handling not specified",
                "description": "Formula can produce negative momentum values but downstream handling is not specified",
                "location": "Section 1 - Momentum Score Calculation",
                "original_text": "Momentum Score = (Current Price / Historical Price) - 1",
                "suggested_fix": "Specify: 'Negative momentum stocks receive zero weight' or 'Use absolute value' or 'Allow negative weights'",
                "confidence": 0.82
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Score normalization not specified",
                "description": "Weight calculation assumes raw scores but doesn't specify if scores should be normalized first",
                "location": "Section 4 - Index Weight Calculation",
                "original_text": "Weight = Score / Sum(Scores)",
                "suggested_fix": "Specify: 'Scores are used as-is (no normalization)' or 'Scores are normalized to [0,1] range before weighting'",
                "confidence": 0.80
            },
            {
                "severity": "medium",
                "category": "incomplete_formula",
                "title": "Rounding/precision not specified",
                "description": "Formulas don't specify decimal precision or rounding rules",
                "location": "All formula sections",
                "original_text": "Various formulas",
                "suggested_fix": "Specify: 'All calculations performed with float64 precision, final weights rounded to 4 decimal places using banker's rounding'",
                "confidence": 0.75
            },
            {
                "severity": "medium",
                "category": "data_source_unspecified",
                "title": "Price data source not specified",
                "description": "Formulas use prices but don't specify data source or handling of missing data",
                "location": "Section 1 - Momentum Score Calculation",
                "original_text": "Current Price / Historical Price",
                "suggested_fix": "Specify: 'Prices from Bloomberg end-of-day feed. If price missing, use last available closing price within 5 trading days'",
                "confidence": 0.78
            }
        ],
        "self_containment_score": 0.35,
        "implementability_score": 0.30
    }
    
    save_document(doc, "doc_04_incomplete_formulas", metadata)


# Document 5: Market Calendar Issues (Critical Priority)
def create_doc_05_market_calendar():
    """Document with ambiguous market calendar specifications"""
    doc = Document()
    
    add_heading(doc, "Global Multi-Asset Tactical Allocation Strategy", 1)
    
    add_heading(doc, "1. Rebalancing Schedule", 2)
    add_paragraph(doc, """
The portfolio is rebalanced monthly, on the last day of each month. 
If the month-end falls on a non-trading day, rebalancing occurs on the next available day.
""")
    
    add_heading(doc, "2. Momentum Calculation", 2)
    add_paragraph(doc, """
Asset momentum is calculated using returns over the previous 60 days. 
For each asset, calculate:
- 1-month return
- 3-month return  
- 6-month return

Composite momentum is the average of these three periods.
""")
    
    add_heading(doc, "3. Correlation Matrix", 2)
    add_paragraph(doc, """
A correlation matrix is computed using the past 90 days of daily returns for all assets. 
This matrix is used to optimize portfolio allocation.
""")
    
    add_heading(doc, "4. Cross-Asset Analysis", 2)
    add_paragraph(doc, """
When analyzing correlations between US equities and Japanese bonds, align the data using 
common trading days to ensure proper comparison.
""")
    
    metadata = {
        "document_id": "doc_05_market_calendar",
        "title": "Global Multi-Asset Tactical Allocation Strategy",
        "version": "1.0",
        "created_date": datetime.now().isoformat(),
        "complexity": "high",
        "issue_count": 11,
        "issues": [
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Market calendar not specified for 'last day of month'",
                "description": "States 'last day of each month' but doesn't specify which market's calendar determines trading days",
                "location": "Section 1 - Rebalancing Schedule",
                "original_text": "last day of each month",
                "suggested_fix": "Specify: 'last NYSE trading day of each calendar month'",
                "confidence": 0.95
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "'Next available day' ambiguous across markets",
                "description": "If month-end is non-trading, uses 'next available day' but doesn't specify which market's next trading day",
                "location": "Section 1 - Rebalancing Schedule",
                "original_text": "next available day",
                "suggested_fix": "Specify: 'next NYSE trading day (even if it falls in the following calendar month)'",
                "confidence": 0.92
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "60 days ambiguous - trading or calendar days?",
                "description": "Momentum uses 'previous 60 days' but doesn't specify if these are trading days or calendar days",
                "location": "Section 2 - Momentum Calculation",
                "original_text": "previous 60 days",
                "suggested_fix": "Specify: 'previous 60 trading days (NYSE calendar)' or 'previous 60 calendar days'",
                "confidence": 0.96
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "1-month period ambiguous",
                "description": "'1-month return' could mean 20 trading days, 30 calendar days, or 1 calendar month",
                "location": "Section 2 - Momentum Calculation",
                "original_text": "1-month return",
                "suggested_fix": "Specify: '1-month return calculated over 21 NYSE trading days'",
                "confidence": 0.94
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "3-month period ambiguous",
                "description": "'3-month return' could mean 60 trading days, 90 calendar days, or 3 calendar months",
                "location": "Section 2 - Momentum Calculation",
                "original_text": "3-month return",
                "suggested_fix": "Specify: '3-month return calculated over 63 NYSE trading days'",
                "confidence": 0.94
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "6-month period ambiguous",
                "description": "'6-month return' could mean 120 trading days, 180 calendar days, or 6 calendar months",
                "location": "Section 2 - Momentum Calculation",
                "original_text": "6-month return",
                "suggested_fix": "Specify: '6-month return calculated over 126 NYSE trading days'",
                "confidence": 0.94
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "90 days for correlation matrix - trading or calendar?",
                "description": "Correlation matrix uses 'past 90 days' without specifying trading vs calendar days",
                "location": "Section 3 - Correlation Matrix",
                "original_text": "past 90 days",
                "suggested_fix": "Specify: 'past 90 calendar days of available daily returns (excluding weekends and holidays)'",
                "confidence": 0.93
            },
            {
                "severity": "critical",
                "category": "market_calendar",
                "title": "Cross-market alignment strategy not specified",
                "description": "US equities and Japanese bonds have different market calendars and holidays. Alignment strategy not specified.",
                "location": "Section 4 - Cross-Asset Analysis",
                "original_text": "align the data using common trading days",
                "suggested_fix": "Specify: 'Use only dates when both NYSE and TSE are open (intersection)' or 'Use all available data with forward-fill for market closures' or 'Align to NYSE calendar, forward-fill Japanese data'",
                "confidence": 0.97
            },
            {
                "severity": "high",
                "category": "ambiguous_methodology",
                "title": "Multi-asset implies multiple calendars not addressed",
                "description": "Strategy is 'Global Multi-Asset' but doesn't specify how to handle different market calendars for different assets",
                "location": "Section 1 - Rebalancing Schedule",
                "original_text": "Global Multi-Asset Tactical Allocation Strategy",
                "suggested_fix": "Add section specifying primary market calendar for rebalancing and how to handle assets from other markets",
                "confidence": 0.88
            },
            {
                "severity": "high",
                "category": "market_calendar",
                "title": "Timezone not specified for rebalancing",
                "description": "Rebalancing occurs 'on the last day' but doesn't specify time or timezone, important for global assets",
                "location": "Section 1 - Rebalancing Schedule",
                "original_text": "rebalanced monthly, on the last day of each month",
                "suggested_fix": "Specify: 'Rebalancing executed at 16:00 America/New_York (NYSE close)'",
                "confidence": 0.85
            },
            {
                "severity": "medium",
                "category": "market_calendar",
                "title": "Holiday handling for different markets not specified",
                "description": "Global strategy but no specification for handling holidays that differ across markets",
                "location": "Section 2 - Momentum Calculation",
                "original_text": "previous 60 days",
                "suggested_fix": "Specify: 'For assets on different exchanges, use asset-specific trading days. Exclude days when that specific market is closed.'",
                "confidence": 0.80
            }
        ],
        "self_containment_score": 0.25,
        "implementability_score": 0.20,
        "notes": "This document exemplifies the critical market calendar issues identified in ADR-018"
    }
    
    save_document(doc, "doc_05_market_calendar", metadata)


# Generate remaining documents...
def main():
    """Generate all test documents"""
    print("Generating test document suite...")
    print(f"Output directory: {OUTPUT_DIR}\n")
    
    # Generate first 5 documents
    create_doc_01_clean()
    create_doc_02_missing_appendix()
    create_doc_03_undefined_parameters()
    create_doc_04_incomplete_formulas()
    create_doc_05_market_calendar()
    
    print(f"\n✓ Successfully created 5 document pairs in {OUTPUT_DIR}")
    print("\nNext: Run 'python scripts/generate_test_documents_part2.py' to create remaining 15 documents")


if __name__ == "__main__":
    main()
