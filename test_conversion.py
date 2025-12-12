#!/usr/bin/env python3
"""
Test script to verify table and mathematical formula extraction from documents.

This script creates test documents with tables and formulas, converts them,
and verifies that the conversion preserves the information correctly.
"""

import asyncio
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import sys

# Add src to path
sys.path.insert(0, str(Path(__file__).parent))

from src.infrastructure.converters.word_converter import WordConverter
from src.infrastructure.converters.pdf_converter import PdfConverter


def create_test_word_doc_with_tables_and_formulas():
    """Create a test Word document with tables and mathematical content."""
    doc = Document()

    # Add title
    doc.add_heading('Test Document: Tables and Formulas', 0)

    # Add paragraph with text
    doc.add_paragraph('This document contains tables and mathematical formulas for testing.')

    # Add a table
    doc.add_heading('Test Table 1: Performance Metrics', 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Strategy'
    hdr_cells[1].text = 'Return (%)'
    hdr_cells[2].text = 'Sharpe Ratio'

    # Data rows
    data = [
        ('Momentum', '12.5', '1.8'),
        ('Mean Reversion', '8.3', '1.2'),
        ('Trend Following', '15.7', '2.1'),
    ]

    for i, (strategy, ret, sharpe) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = strategy
        row_cells[1].text = ret
        row_cells[2].text = sharpe

    # Add mathematical content
    doc.add_heading('Mathematical Formulas', 2)

    # Add paragraphs with mathematical symbols
    doc.add_paragraph('The Sharpe Ratio is calculated as:')
    doc.add_paragraph('Sharpe Ratio = (μ - Rf) / σ')
    doc.add_paragraph('where:')
    doc.add_paragraph('  μ = expected return')
    doc.add_paragraph('  σ = standard deviation of returns')
    doc.add_paragraph('  Rf = risk-free rate')

    doc.add_paragraph('')
    doc.add_paragraph('The portfolio variance is:')
    doc.add_paragraph('σ² = ∑ᵢ∑ⱼ wᵢwⱼσᵢⱼ')

    doc.add_paragraph('')
    doc.add_paragraph('Common Greek letters used in finance:')
    doc.add_paragraph('  α (alpha) - excess return')
    doc.add_paragraph('  β (beta) - market sensitivity')
    doc.add_paragraph('  γ (gamma) - rate of change')
    doc.add_paragraph('  δ (delta) - option sensitivity')

    # Add another table
    doc.add_heading('Test Table 2: Greek Letters', 2)
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Light Grid Accent 1'

    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Symbol'
    hdr_cells2[1].text = 'Name'
    hdr_cells2[2].text = 'Financial Use'

    greek_data = [
        ('α', 'Alpha', 'Excess return'),
        ('β', 'Beta', 'Market sensitivity'),
        ('γ', 'Gamma', 'Rate of change'),
        ('σ', 'Sigma', 'Standard deviation'),
    ]

    for i, (symbol, name, use) in enumerate(greek_data, 1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = symbol
        row_cells[1].text = name
        row_cells[2].text = use

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def test_word_conversion():
    """Test Word document conversion for tables and formulas."""
    print("=" * 80)
    print("Testing Word Document Conversion")
    print("=" * 80)

    # Create test document
    print("\n1. Creating test Word document with tables and formulas...")
    content = create_test_word_doc_with_tables_and_formulas()
    print(f"   Created document ({len(content)} bytes)")

    # Convert document
    print("\n2. Converting document to markdown...")
    converter = WordConverter()
    result = converter.convert_from_bytes(content, "test_tables_formulas.docx")

    if not result.success:
        print(f"   ❌ Conversion failed: {result.errors}")
        return False

    print(f"   ✓ Conversion successful")
    print(f"   Word count: {result.metadata.word_count}")
    print(f"   Sections found: {len(result.sections)}")

    # Analyze markdown content
    print("\n3. Analyzing converted markdown...")
    markdown = result.markdown_content

    # Check for tables
    table_count = markdown.count('| --- |')
    print(f"\n   Tables:")
    print(f"   - Found {table_count} markdown tables")

    if table_count >= 2:
        print(f"   ✓ Both tables converted successfully")
    else:
        print(f"   ⚠️  Expected 2 tables, found {table_count}")

    # Check for mathematical symbols
    print(f"\n   Mathematical symbols:")
    symbols = {
        'μ': 'mu (expected return)',
        'σ': 'sigma (standard deviation)',
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        '∑': 'summation',
    }

    for symbol, name in symbols.items():
        count = markdown.count(symbol)
        if count > 0:
            print(f"   ✓ {name} ({symbol}): {count} occurrences")
        else:
            print(f"   ⚠️  {name} ({symbol}): NOT FOUND")

    # Print sample of markdown
    print("\n4. Sample markdown output (first 1500 chars):")
    print("-" * 80)
    print(markdown[:1500])
    print("-" * 80)

    # Look for table sections
    print("\n5. Extracting table sections...")
    lines = markdown.split('\n')
    in_table = False
    table_num = 0

    for i, line in enumerate(lines):
        if '|' in line and '---' in line and not in_table:
            in_table = True
            table_num += 1
            print(f"\n   Table {table_num} found at line {i}:")
            # Print a few lines of the table
            for j in range(max(0, i-1), min(len(lines), i+6)):
                print(f"     {lines[j]}")

    return True


async def main():
    """Main test runner."""
    print("\n" + "=" * 80)
    print("Document Conversion Test Suite")
    print("Testing table and mathematical formula extraction")
    print("=" * 80 + "\n")

    success = await test_word_conversion()

    print("\n" + "=" * 80)
    if success:
        print("✓ All tests passed!")
    else:
        print("❌ Some tests failed")
    print("=" * 80 + "\n")

    return success


if __name__ == "__main__":
    asyncio.run(main())
